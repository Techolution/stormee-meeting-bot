import logging
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from utilities.logging_context import RequestContext

logger = logging.getLogger(__name__)


async def save_captions_to_docx(filename: Optional[str] = None, captions_segments: list = None, title: Optional[str] = None) -> str:
    """Exports all gathered meeting captions into a formatted DOCX document.
    
    Args:
        filename: Optional output filename. If not provided, auto-generated with timestamp
        captions_segments: List of caption entries with speaker, text, timestamp
        title: Optional document title for the DOCX file
        
    Returns:
        Path to the saved DOCX file, or empty string if no captions to export
        
    Raises:
        Exception: If document generation or save fails
    """
    context = RequestContext.get_context()
    
    logger.debug(
        f"Initiating caption export to DOCX [segments_count={len(captions_segments) if captions_segments else 0}, "
        f"title={title}, request_id={context.get('request_id', 'N/A')}]"
    )
    
    if not captions_segments:
        logger.warning(
            f"No caption segments available to export [title={title}, "
            f"request_id={context.get('request_id', 'N/A')}]"
        )
        return ""

    if not filename:
        timestamp_str = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"meeting_captions_{timestamp_str}.docx"
        logger.debug(f"Generated filename for DOCX export: {filename}")

    try:
        doc = Document()
        doc_title = f"{title} Transcript" if title else "Google Meet Transcript"
        # Title
        title_heading = doc.add_heading(doc_title, level=0)
        title_heading.style.font.name = "Arial"

        # Metadata Paragraph
        meta = doc.add_paragraph()
        meta.add_run(
            f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        ).italic = True
        meta.add_run(
            f"Total Caption Blocks: {len(captions_segments)}"
        ).italic = True

        doc.add_heading("Transcript", level=1)

        # Render each caption turn
        processed_count = 0
        for entry in captions_segments:
            speaker = entry.get("speaker", "Unknown Speaker")
            text = entry.get("text", "")
            time_str = entry.get("timestamp", "")

            try:
                parsed_time = datetime.fromisoformat(time_str).strftime("%H:%M:%S")
            except Exception as e:
                logger.debug(f"Failed to parse timestamp '{time_str}': {e}, using raw value")
                parsed_time = time_str

            p = doc.add_paragraph()

            # Speaker & Time
            speaker_run = p.add_run(f"[{parsed_time}] {speaker}: ")
            speaker_run.bold = True
            speaker_run.font.color.rgb = RGBColor(0, 102, 204)  # Soft blue

            # Caption Body
            p.add_run(text)
            processed_count += 1

        # Save document
        doc.save(filename)
        
        logger.info(
            f"Captions successfully exported to DOCX [file={filename}, "
            f"segments_processed={processed_count}, "
            f"request_id={context.get('request_id', 'N/A')}]"
        )
        
        return filename
        
    except Exception as e:
        logger.error(
            f"Failed to save captions to DOCX [filename={filename}, "
            f"segments_count={len(captions_segments)}, "
            f"error_type={type(e).__name__}, "
            f"error_msg={str(e)[:200]}, "
            f"request_id={context.get('request_id', 'N/A')}]"
        )
        raise