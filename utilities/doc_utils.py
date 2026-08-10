from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor


async def save_captions_to_docx(filename: Optional[str] = None, captions_segments: list = None, title: Optional[str] = None) -> str:
    """Exports all gathered meeting captions into a formatted DOCX document."""
    if not captions_segments:
        print("⚠️ No caption segments available to export.")
        return ""

    if not filename:
        timestamp_str = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"meeting_captions_{timestamp_str}.docx"

    doc = Document()
    doc_title = f"{title} Transcript" if title else "Google Meet Transcript"
    # Title
    title = doc.add_heading(doc_title, level=0)
    title.style.font.name = "Arial"

    # Metadata Paragraph
    meta = doc.add_paragraph()
    meta.add_run(
        f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
    ).italic = True
    meta.add_run(
        f"Total Caption Blocks: {len(captions_segments)}"
    ).italic = True

    doc.add_heading("Transcript", level=1)

    # Render each caption turn
    for entry in captions_segments:
        speaker = entry.get("speaker", "Unknown Speaker")
        text = entry.get("text", "")
        time_str = entry.get("timestamp", "")

        try:
            parsed_time = datetime.fromisoformat(time_str).strftime("%H:%M:%S")
        except Exception:
            parsed_time = time_str

        p = doc.add_paragraph()

        # Speaker & Time
        speaker_run = p.add_run(f"[{parsed_time}] {speaker}: ")
        speaker_run.bold = True
        speaker_run.font.color.rgb = RGBColor(0, 102, 204)  # Soft blue

        # Caption Body
        p.add_run(text)

    doc.save(filename)
    print(f"📄 Captions successfully exported to Word Document: {filename}")
    return filename